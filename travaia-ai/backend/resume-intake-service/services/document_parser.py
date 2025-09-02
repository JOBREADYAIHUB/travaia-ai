"""
Document Parser Service
Handles parsing of resume documents (PDF, DOCX) and text extraction
"""

import os
import re
from typing import Optional, Dict, Any, List
from datetime import datetime
import asyncio
import logging
from pydantic import BaseModel

# PDF parsing
try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    PyPDF2 = None

# DOCX parsing
try:
    from docx import Document
except ImportError:
    Document = None

# Text processing
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

logger = logging.getLogger(__name__)

class ParsedResumeData(BaseModel):
    personalInfo: Dict[str, Any]
    summary: Optional[str] = None
    experience: List[Dict[str, Any]] = []
    education: List[Dict[str, Any]] = []
    skills: List[Dict[str, Any]] = []
    certifications: List[Dict[str, Any]] = []
    rawText: str
    confidence: float

class DocumentParser:
    """Service for parsing resume documents and extracting structured data"""
    
    def __init__(self):
        self._ensure_nltk_data()
        
    def _ensure_nltk_data(self):
        """Ensure required NLTK data is available"""
        try:
            nltk.data.find('tokenizers/punkt')
            nltk.data.find('corpora/stopwords')
        except LookupError:
            logger.info("Downloading required NLTK data...")
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
    
    async def parse_document(
        self, 
        file_content: bytes, 
        file_type: str, 
        filename: str
    ) -> ParsedResumeData:
        """Parse document and extract structured resume data"""
        try:
            # Extract raw text based on file type
            if file_type.lower() == '.pdf':
                raw_text = await self._parse_pdf(file_content)
            elif file_type.lower() in ['.docx', '.doc']:
                raw_text = await self._parse_docx(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            if not raw_text or len(raw_text.strip()) < 50:
                raise ValueError("Insufficient text content extracted from document")
            
            # Extract structured data from raw text
            parsed_data = await self._extract_structured_data(raw_text)
            
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error parsing document {filename}: {str(e)}")
            # Return minimal data structure with raw text
            return ParsedResumeData(
                personalInfo={},
                rawText=raw_text if 'raw_text' in locals() else "",
                confidence=0.0
            )
    
    async def _parse_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        if not PyPDF2:
            raise ImportError("PyPDF2 not available for PDF parsing")
        
        try:
            from io import BytesIO
            pdf_file = BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            
            text_content = []
            for page in pdf_reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting text from PDF page: {str(e)}")
                    continue
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise Exception(f"Failed to parse PDF document: {str(e)}")
    
    async def _parse_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        if not Document:
            raise ImportError("python-docx not available for DOCX parsing")
        
        try:
            from io import BytesIO
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise Exception(f"Failed to parse DOCX document: {str(e)}")
    
    async def _extract_structured_data(self, raw_text: str) -> ParsedResumeData:
        """Extract structured data from raw text"""
        try:
            # Clean and normalize text
            cleaned_text = self._clean_text(raw_text)
            
            # Extract different sections
            personal_info = self._extract_personal_info(cleaned_text)
            summary = self._extract_summary(cleaned_text)
            experience = self._extract_experience(cleaned_text)
            education = self._extract_education(cleaned_text)
            skills = self._extract_skills(cleaned_text)
            certifications = self._extract_certifications(cleaned_text)
            
            # Calculate confidence score
            confidence = self._calculate_confidence(
                personal_info, summary, experience, education, skills
            )
            
            return ParsedResumeData(
                personalInfo=personal_info,
                summary=summary,
                experience=experience,
                education=education,
                skills=skills,
                certifications=certifications,
                rawText=raw_text,
                confidence=confidence
            )
            
        except Exception as e:
            logger.error(f"Error extracting structured data: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\-\.\@\(\)\,\:\;]', ' ', text)
        return text.strip()
    
    def _extract_personal_info(self, text: str) -> Dict[str, Any]:
        """Extract personal information"""
        personal_info = {}
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            personal_info['email'] = emails[0]
        
        # Extract phone number
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        if phones:
            personal_info['phone'] = ''.join(phones[0]) if isinstance(phones[0], tuple) else phones[0]
        
        # Extract name (first few words before email or phone)
        lines = text.split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if line and len(line.split()) <= 4 and not any(char.isdigit() for char in line):
                if '@' not in line and not re.search(r'\d{3}', line):
                    personal_info['name'] = line
                    break
        
        # Extract location/address
        location_keywords = ['address', 'location', 'city', 'state', 'country']
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in location_keywords):
                # Extract location from this line
                location_match = re.search(r'([A-Za-z\s]+,\s*[A-Za-z\s]+)', line)
                if location_match:
                    personal_info['location'] = location_match.group(1)
                    break
        
        return personal_info
    
    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract professional summary or objective"""
        summary_keywords = [
            'summary', 'objective', 'profile', 'about', 'overview',
            'professional summary', 'career objective', 'personal statement'
        ]
        
        lines = text.split('\n')
        summary_start = -1
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in summary_keywords):
                summary_start = i
                break
        
        if summary_start >= 0:
            # Extract next few lines as summary
            summary_lines = []
            for i in range(summary_start + 1, min(summary_start + 6, len(lines))):
                line = lines[i].strip()
                if line and not self._is_section_header(line):
                    summary_lines.append(line)
                elif line and self._is_section_header(line):
                    break
            
            if summary_lines:
                return ' '.join(summary_lines)
        
        return None
    
    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience"""
        experience_keywords = [
            'experience', 'employment', 'work history', 'professional experience',
            'career history', 'work experience'
        ]
        
        lines = text.split('\n')
        experience_start = -1
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in experience_keywords):
                experience_start = i
                break
        
        experiences = []
        if experience_start >= 0:
            current_job = {}
            
            for i in range(experience_start + 1, len(lines)):
                line = lines[i].strip()
                if not line:
                    continue
                
                # Check if this is a new section
                if self._is_section_header(line) and not any(keyword in line.lower() for keyword in experience_keywords):
                    break
                
                # Try to identify job title, company, dates
                if self._looks_like_job_title(line):
                    if current_job:
                        experiences.append(current_job)
                    current_job = {'title': line, 'description': []}
                elif self._looks_like_company(line):
                    if current_job:
                        current_job['company'] = line
                elif self._looks_like_dates(line):
                    if current_job:
                        current_job['dates'] = line
                else:
                    if current_job:
                        current_job['description'].append(line)
            
            if current_job:
                experiences.append(current_job)
        
        # Clean up descriptions
        for exp in experiences:
            if 'description' in exp:
                exp['description'] = ' '.join(exp['description'])
        
        return experiences
    
    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education information"""
        education_keywords = ['education', 'academic', 'degree', 'university', 'college']
        
        lines = text.split('\n')
        education_start = -1
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in education_keywords):
                education_start = i
                break
        
        education_list = []
        if education_start >= 0:
            current_edu = {}
            
            for i in range(education_start + 1, len(lines)):
                line = lines[i].strip()
                if not line:
                    continue
                
                if self._is_section_header(line) and not any(keyword in line.lower() for keyword in education_keywords):
                    break
                
                # Try to identify degree, institution, dates
                if self._looks_like_degree(line):
                    if current_edu:
                        education_list.append(current_edu)
                    current_edu = {'degree': line}
                elif self._looks_like_institution(line):
                    if current_edu:
                        current_edu['institution'] = line
                elif self._looks_like_dates(line):
                    if current_edu:
                        current_edu['dates'] = line
            
            if current_edu:
                education_list.append(current_edu)
        
        return education_list
    
    def _extract_skills(self, text: str) -> List[Dict[str, Any]]:
        """Extract skills information"""
        skills_keywords = ['skills', 'technical skills', 'competencies', 'technologies']
        
        lines = text.split('\n')
        skills_start = -1
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in skills_keywords):
                skills_start = i
                break
        
        skills_list = []
        if skills_start >= 0:
            for i in range(skills_start + 1, len(lines)):
                line = lines[i].strip()
                if not line:
                    continue
                
                if self._is_section_header(line):
                    break
                
                # Split skills by common delimiters
                skill_items = re.split(r'[,;|•·]', line)
                for skill in skill_items:
                    skill = skill.strip()
                    if skill and len(skill) > 1:
                        skills_list.append({
                            'name': skill,
                            'category': 'general'
                        })
        
        return skills_list
    
    def _extract_certifications(self, text: str) -> List[Dict[str, Any]]:
        """Extract certifications"""
        cert_keywords = ['certification', 'certificate', 'licensed', 'certified']
        
        lines = text.split('\n')
        certifications = []
        
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in cert_keywords):
                certifications.append({
                    'name': line.strip(),
                    'issuer': 'Unknown'
                })
        
        return certifications
    
    def _is_section_header(self, line: str) -> bool:
        """Check if line is likely a section header"""
        line = line.strip()
        if len(line) < 3:
            return False
        
        section_keywords = [
            'experience', 'education', 'skills', 'summary', 'objective',
            'certifications', 'projects', 'achievements', 'awards'
        ]
        
        return any(keyword in line.lower() for keyword in section_keywords)
    
    def _looks_like_job_title(self, line: str) -> bool:
        """Check if line looks like a job title"""
        job_keywords = ['manager', 'developer', 'engineer', 'analyst', 'specialist', 'coordinator']
        return any(keyword in line.lower() for keyword in job_keywords)
    
    def _looks_like_company(self, line: str) -> bool:
        """Check if line looks like a company name"""
        company_indicators = ['inc', 'llc', 'corp', 'ltd', 'company', 'technologies']
        return any(indicator in line.lower() for indicator in company_indicators)
    
    def _looks_like_dates(self, line: str) -> bool:
        """Check if line contains dates"""
        date_pattern = r'\b\d{4}\b|\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)'
        return bool(re.search(date_pattern, line.lower()))
    
    def _looks_like_degree(self, line: str) -> bool:
        """Check if line looks like a degree"""
        degree_keywords = ['bachelor', 'master', 'phd', 'degree', 'diploma', 'certificate']
        return any(keyword in line.lower() for keyword in degree_keywords)
    
    def _looks_like_institution(self, line: str) -> bool:
        """Check if line looks like an educational institution"""
        institution_keywords = ['university', 'college', 'institute', 'school']
        return any(keyword in line.lower() for keyword in institution_keywords)
    
    def _calculate_confidence(
        self, 
        personal_info: Dict, 
        summary: Optional[str], 
        experience: List, 
        education: List, 
        skills: List
    ) -> float:
        """Calculate confidence score for parsed data"""
        score = 0.0
        
        # Personal info scoring
        if personal_info.get('name'):
            score += 0.2
        if personal_info.get('email'):
            score += 0.2
        if personal_info.get('phone'):
            score += 0.1
        
        # Content scoring
        if summary:
            score += 0.1
        if experience:
            score += 0.2
        if education:
            score += 0.1
        if skills:
            score += 0.1
        
        return min(score, 1.0)
